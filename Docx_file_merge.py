from docx import Document
import os

def merge_docx_files(output_filename, directory):
    """
    合并同一目录下的所有.docx文件到一个文档中，并在每个文件内容前加上文件名。
    
    :param output_filename: 输出文件名（含路径）
    :param directory: 包含.docx文件的目录
    """
    # 创建一个新的Document对象作为输出文档
    merged_document = Document()

    # 遍历指定目录下的所有.docx文件
    for filename in os.listdir(directory):
        if filename.endswith('.docx'):
            file_path = os.path.join(directory, filename)
            print(f"正在处理文件: {file_path}")
            
            # 读取当前文件
            current_document = Document(file_path)
            
            # 在当前文件内容前加上文件名，并设置为标题格式
            heading_paragraph = merged_document.add_paragraph()
            heading_run = heading_paragraph.add_run(f'{filename}\n\n')
            heading_paragraph.style = 'Heading 1'  # 设置为标题样式
            
            # 将当前文件的内容逐段追加到输出文档
            for paragraph in current_document.paragraphs:
                new_paragraph = merged_document.add_paragraph(paragraph.text)
                # 复制段落格式
                new_paragraph.style = current_document.styles[paragraph.style.name]
                
                # 复制段落内的Run对象
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text or '')
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    
            # 添加一个分页符以区分不同文件
            merged_document.add_page_break()

    # 保存合并后的文档
    output_path = os.path.join(directory, output_filename)
    print(f"尝试保存到: {output_path}")
    try:
        merged_document.save(output_path)
        print(f"合并完成，已保存为 {output_filename}")
    except Exception as e:
        print(f"保存文件时发生错误: {e}")

# 定义输入和输出目录及文件名
input_directory = r'C:\Users\Yi\Desktop\scraper'
output_filename = 'merged.docx'

# 调用函数合并文件
merge_docx_files(output_filename, input_directory)